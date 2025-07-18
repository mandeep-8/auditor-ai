#pages/cdd_processing.py
import streamlit as st
import pandas as pd
import os
from datetime import datetime
from utils.question_answering import get_qa_chain
from utils.session_logger import get_session_logger
from docx import Document
import base64
import time
from config import INPUT_DIR, OUTPUT_DIR, LOG_DIR, SESSIONS_DIR, UPLOAD_DIR

def save_docx(doc, filename, session_logger):
    """Save DOCX to outputs directory."""
    docx_path = os.path.join(OUTPUT_DIR, filename)
    try:
        doc.save(docx_path)
        session_logger.log(
            component="CDD Report",
            message=f"Saving DOCX: {filename}",
            decision="Accepted",
            reason="DOCX saved successfully",
            level="INFO",
            context={"file": docx_path}
        )
        return docx_path
    except (IOError, PermissionError) as e:
        session_logger.log(
            component="CDD Report",
            message=f"Saving DOCX: {filename}",
            decision="Rejected",
            reason=f"Failed to save: {str(e)}",
            level="ERROR",
            context={"file": docx_path}
        )
        raise

def get_docx_download_link(docx_path, link_text):
    """Generate a download link for a DOCX file."""
    with open(docx_path, "rb") as f:
        docx_data = f.read()
    b64 = base64.b64encode(docx_data).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{os.path.basename(docx_path)}">{link_text}</a>'

def validate_approver_titles(df, question_col, answer_col, qa_chain, session_logger):
    """Validate approver job titles in batch using the QA chain."""
    if qa_chain is None:
        session_logger.log(
            component="CDD Processing",
            message="Approver Title Validation",
            decision="Skipped",
            reason="QA chain not initialized due to missing Governance Documents",
            level="INFO",
            context={}
        )
        return []

    missing_titles = []
    approver_col = next((col for col in df.columns if col.lower().replace(" ", "") == "approvername"), None)
    if not approver_col:
        session_logger.log(
            component="CDD Processing",
            message="Approver Title Validation",
            decision="Rejected",
            reason="Approver Name column not found in Excel",
            level="ERROR",
            context={"columns": list(df.columns)}
        )
        return missing_titles

    # Collect queries for batch processing
    queries = []
    query_indices = []
    for idx, row in df.iterrows():
        question = row[question_col]
        answer = row[answer_col] if answer_col in row and pd.notna(row[answer_col]) else ""
        approver_name = row[approver_col] if pd.notna(row[approver_col]) else None
        if pd.isna(question) or not approver_name:
            continue
        queries.append(f"What is the job title of {approver_name} in this answer?")
        query_indices.append((idx, question, approver_name, answer))

    if queries:
        try:
            # Batch invoke LLM
            batch_prompt = "\n".join([f"Query {i+1}: {q}" for i, q in enumerate(queries)])
            batch_response = qa_chain.invoke({
                "query": f"Answer the following queries in a JSON list of objects with 'query_index' and 'job_title':\n{batch_prompt}",
                "context": " ".join([row[answer_col] for _, row in df.iterrows() if answer_col in row and pd.notna(row[answer_col])])
            })
            responses = eval(batch_response["result"]) if isinstance(batch_response["result"], str) else batch_response["result"]

            for (idx, question, approver_name, answer), response in zip(query_indices, responses):
                job_title = response.get("job_title", "").strip()
                if not job_title or job_title.lower() in ["unknown", "not found", "none"]:
                    missing_titles.append({"index": idx, "question": question, "approver_name": approver_name})
                    session_logger.log(
                        component="CDD Processing",
                        message=f"Row {idx + 1}: Question={question}",
                        decision="Pending",
                        reason=f"Job title for {approver_name} not found in answer",
                        level="WARNING",
                        context={"answer": answer[:50] + "..."}
                    )
                else:
                    session_logger.log(
                        component="CDD Processing",
                        message=f"Row {idx + 1}: Question={question}",
                        decision="Accepted",
                        reason=f"Job title for {approver_name}: {job_title}",
                        level="INFO",
                        context={"answer": answer[:50] + "..."}
                    )
        except Exception as e:
            for idx, question, approver_name, _ in query_indices:
                missing_titles.append({"index": idx, "question": question, "approver_name": approver_name})
                session_logger.log(
                    component="CDD Processing",
                    message=f"Row {idx + 1}: Question={question}",
                    decision="Error",
                    reason=f"Title validation error: {str(e)}",
                    level="ERROR",
                    context={"approver_name": approver_name}
                )

    return missing_titles

def highlight_missing_titles(df, missing_titles):
    """Apply dark blue background to rows with missing approver titles."""
    def highlight_row(row):
        if row.name in [m['index'] for m in missing_titles]:
            return ['background-color: #00008B; color: white'] * len(row)
        return [''] * len(row)
    return df.style.apply(highlight_row, axis=1)

def process_cdd():
    """Process Control Design Document and update Excel in UI."""
    session_logger = get_session_logger(LOG_DIR, st.session_state.session_id)
    st.subheader("📊 Control Design Document Processing")
    session_logger.log(
        component="CDD Processing",
        message="Started CDD processing",
        decision="Accepted",
        level="INFO",
        context={}
    )

    # Initialize QA chain
    try:
        qa_chain_result = get_qa_chain()
        st.session_state.qa_chain = qa_chain_result
        session_logger.log(
            component="CDD Processing",
            message="QA Chain Initialization",
            decision="Accepted",
            reason="QA chain initialized successfully",
            level="INFO",
            context={}
        )
    except Exception as e:
        st.session_state.qa_chain = None
        session_logger.log(
            component="CDD Processing",
            message="QA Chain Initialization",
            decision="Skipped",
            reason=f"No valid Governance Documents to index: {str(e)}",
            level="WARNING",
            context={}
        )

    qa_chain = st.session_state.get("qa_chain", None)

    # Initialize session state
    if "pip_control_owner_name" not in st.session_state:
        st.session_state.pip_control_owner_name = ""
    if "pip_control_owner_title" not in st.session_state:
        st.session_state.pip_control_owner_title = ""
    if "cdd_saved" not in st.session_state:
        st.session_state.cdd_saved = False
    if "cdd_validated" not in st.session_state:
        st.session_state.cdd_validated = False
    if "cdd_text" not in st.session_state:
        st.session_state.cdd_text = ""
    if "excel_df" not in st.session_state:
        st.session_state.excel_df = None
    if "edit_mode" not in st.session_state:
        st.session_state.edit_mode = False
    if "questions_processed" not in st.session_state:
        st.session_state.questions_processed = False

    # Sidebar for Control Owner Profile
    with st.sidebar:
        st.markdown("### Control Owner Profile")
        if st.session_state.pip_control_owner_name and st.session_state.pip_control_owner_title:
            st.markdown(
                f"""
                <div style="border: 2px solid #4CAF50; padding: 10px; border-radius: 5px; margin: 10px 0;">
                    <strong>Name:</strong> {st.session_state.pip_control_owner_name}<br>
                    <strong>Designation:</strong> {st.session_state.pip_control_owner_title}
                </div>
                """,
                unsafe_allow_html=True
            )
        else:
            st.warning("⚠️ Please submit the PIP form to provide Control Owner details.")
            session_logger.log(
                component="CDD Processing",
                message="Control Owner Profile",
                decision="Pending",
                reason="Control Owner details not found",
                level="WARNING",
                context={}
            )

    # Validate Control Owner details
    if not (st.session_state.pip_control_owner_name and st.session_state.pip_control_owner_title):
        st.error("⚠️ Please submit the PIP form to provide Control Owner Name and Designation before proceeding.")
        session_logger.log(
            component="CDD Processing",
            message="Control Owner Validation",
            decision="Rejected",
            reason="Control Owner Name or Designation missing",
            level="ERROR",
            context={}
        )
        return

    excel_path = os.path.join(INPUT_DIR, "Control Design Document.xlsx")
    if not os.path.exists(excel_path):
        st.error("⚠️ Control Design Document.xlsx not found in inputs directory.")
        session_logger.log(
            component="CDD Processing",
            message="Excel File Check",
            decision="Rejected",
            reason="Control Design Document.xlsx not found",
            level="ERROR",
            context={"file": excel_path}
        )
        return

    try:
        df = pd.read_excel(excel_path, engine="openpyxl", dtype=str)
        actual_columns = {col.lower().replace(" ", ""): col for col in df.columns}
        question_col = next((col for key, col in actual_columns.items() if "controldesignunderstandingquestions" in key), None)
        answer_col = next((col for key, col in actual_columns.items() if "answers" in key), None)
        for col in df.columns:
            if col not in [question_col, answer_col]:
                try:
                    df[col] = pd.to_numeric(df[col], errors='coerce')
                    if df[col].isna().all():
                        df[col] = df[col].astype("string")
                except:
                    df[col] = df[col].astype("string")
        if question_col:
            df[question_col] = df[question_col].astype("string")
        if answer_col:
            df[answer_col] = df[answer_col].astype("string")
        session_logger.log(
            component="CDD Processing",
            message="Excel File Load",
            decision="Accepted",
            reason="Excel file loaded successfully",
            level="INFO",
            context={"columns": list(df.columns)}
        )
    except (FileNotFoundError, ValueError, PermissionError) as e:
        st.error(f"⚠️ Failed to read Excel file: {str(e)}")
        session_logger.log(
            component="CDD Processing",
            message="Excel File Load",
            decision="Rejected",
            reason=f"Failed to read Excel: {str(e)}",
            level="ERROR",
            context={"file": excel_path}
        )
        return

    # Check for required columns
    if not question_col or not answer_col:
        st.error(f"⚠️ Excel file must contain 'Control design understanding questions' and 'Answers' columns. Found: {list(df.columns)}")
        session_logger.log(
            component="CDD Processing",
            message="Column Check",
            decision="Rejected",
            reason=f"Required columns not found: {list(df.columns)}",
            level="ERROR",
            context={"columns": list(df.columns)}
        )
        return

    # Initialize session state for DataFrame
    if st.session_state.excel_df is None:
        st.session_state.excel_df = df.copy()
        session_logger.log(
            component="CDD Processing",
            message="Session State Initialization",
            decision="Accepted",
            reason="Excel DataFrame initialized",
            level="INFO",
            context={}
        )

    # Display Excel in UI
    st.write("### Control Design Document")
    excel_placeholder = st.empty()

    # Toggle edit mode
    if st.button("✏️ Edit Excel", key="edit_excel"):
        st.session_state.edit_mode = not st.session_state.edit_mode
        session_logger.log(
            component="CDD Processing",
            message=f"Edit Mode Toggled: {'Enabled' if st.session_state.edit_mode else 'Disabled'}",
            decision="Accepted",
            reason="Edit mode toggled successfully",
            level="INFO",
            context={}
        )

    # Configure column settings
    column_config = {}
    for col in st.session_state.excel_df.columns:
        dtype = st.session_state.excel_df[col].dtype
        if col == question_col:
            column_config[col] = st.column_config.TextColumn(disabled=True)
        elif col == answer_col:
            column_config[col] = st.column_config.TextColumn(disabled=False)
        elif pd.api.types.is_numeric_dtype(dtype) and not st.session_state.excel_df[col].isna().all():
            if pd.api.types.is_integer_dtype(dtype):
                column_config[col] = st.column_config.NumberColumn(disabled=False, format="%d")
            else:
                column_config[col] = st.column_config.NumberColumn(disabled=False, format="%.2f")
        else:
            column_config[col] = st.column_config.TextColumn(disabled=False)

    # Validate approver titles for highlighting
    missing_titles = validate_approver_titles(st.session_state.excel_df, question_col, answer_col, qa_chain, session_logger)
    if qa_chain is None and not missing_titles:
        st.session_state.cdd_validated = True  # Validate when no Governance Documents and no approver titles required
        session_logger.log(
            component="CDD Processing",
            message="Approver Title Validation",
            decision="Accepted",
            reason="No Governance Documents provided, no approver titles required",
            level="INFO",
            context={}
        )

    if st.session_state.edit_mode:
        edited_df = excel_placeholder.data_editor(
            st.session_state.excel_df,
            use_container_width=True,
            column_config=column_config,
            key="cdd_data_editor"
        )
        st.session_state.excel_df = edited_df
        if st.button("💾 Save Changes", key="submit_changes"):
            session_dir = os.path.join(SESSIONS_DIR, st.session_state.session_id)
            os.makedirs(session_dir, exist_ok=True)
            updated_excel_path = os.path.join(session_dir, "Control Design Document_updated.xlsx")
            try:
                st.session_state.excel_df.to_excel(updated_excel_path, engine="openpyxl", index=False)
                st.success(f"✅ Changes saved to {updated_excel_path}")
                st.session_state.cdd_saved = True
                missing_titles = validate_approver_titles(st.session_state.excel_df, question_col, answer_col, qa_chain, session_logger)
                if not missing_titles:
                    st.session_state.cdd_validated = True
                    session_logger.log(
                        component="CDD Processing",
                        message="Approver Title Validation",
                        decision="Accepted",
                        reason="No missing approver titles after save",
                        level="INFO",
                        context={}
                    )
                answers = st.session_state.excel_df[answer_col].dropna().astype(str)
                st.session_state.cdd_text = " ".join(answers)
                session_logger.log(
                    component="CDD Processing",
                    message="Excel Save",
                    decision="Accepted",
                    reason="Edited Excel saved successfully",
                    level="INFO",
                    context={"file": updated_excel_path}
                )
            except (IOError, PermissionError) as e:
                st.error(f"⚠️ Failed to save Excel: {str(e)}")
                session_logger.log(
                    component="CDD Processing",
                    message="Excel Save",
                    decision="Rejected",
                    reason=f"Failed to save Excel: {str(e)}",
                    level="ERROR",
                    context={"file": updated_excel_path}
                )
    else:
        styled_df = highlight_missing_titles(st.session_state.excel_df, missing_titles)
        excel_placeholder.dataframe(
            styled_df,
            use_container_width=True
        )

    # Process questions and validate approver titles
    if not st.session_state.questions_processed and qa_chain is not None:
        question_placeholder = st.empty()
        invalid_answers = []
        for idx, row in st.session_state.excel_df.iterrows():
            question = row[question_col]
            if pd.isna(question):
                session_logger.log(
                    component="CDD Processing",
                    message=f"Row {idx + 1}",
                    decision="Skipped",
                    reason="Empty question",
                    level="INFO",
                    context={}
                )
                continue

            question_placeholder.markdown(
                f'<div style="border: 2px solid #00008B; padding: 10px; border-radius: 5px; margin: 10px 0;">Row {idx + 1}: {question}</div>',
                unsafe_allow_html=True
            )

            try:
                response = qa_chain.invoke({"query": question})
                qa_answer = response["result"].strip()
                sources = [doc.metadata["source"] for doc in response["source_documents"]]
                if not qa_answer or qa_answer.lower() in ["i don't know", "unknown"]:
                    invalid_answers.append(question)
                    st.session_state.excel_df.at[idx, answer_col] = ""
                    session_logger.log(
                        component="CDD Processing",
                        message=f"Row {idx + 1}: Question={question}",
                        decision="Invalid",
                        reason="QA returned empty or unknown answer",
                        level="WARNING",
                        context={"sources": ', '.join(sources)}
                    )
                else:
                    st.session_state.excel_df.at[idx, answer_col] = qa_answer
                    session_logger.log(
                        component="CDD Processing",
                        message=f"Row {idx + 1}: Question={question}",
                        decision="Accepted",
                        reason=f"Answer={qa_answer}",
                        level="INFO",
                        context={"sources": ', '.join(sources)}
                    )
                styled_df = highlight_missing_titles(st.session_state.excel_df, missing_titles)
                excel_placeholder.dataframe(
                    styled_df,
                    use_container_width=True
                )
            except (ValueError, RuntimeError, AttributeError) as e:
                invalid_answers.append(question)
                st.session_state.excel_df.at[idx, answer_col] = ""
                session_logger.log(
                    component="CDD Processing",
                    message=f"Row {idx + 1}: Question={question}",
                    decision="Error",
                    reason=f"QA Processing Error: {str(e)}",
                    level="ERROR",
                    context={}
                )
                styled_df = highlight_missing_titles(st.session_state.excel_df, missing_titles)
                excel_placeholder.dataframe(
                    styled_df,
                    use_container_width=True
                )

            time.sleep(0.2)
            question_placeholder.empty()

        answers = st.session_state.excel_df[answer_col].dropna().astype(str)
        st.session_state.cdd_text = " ".join(answers)
        session_logger.log(
            component="CDD Processing",
            message="CDD Text Generation",
            decision="Accepted",
            reason="CDD text generated from Excel answers",
            level="INFO",
            context={"text_length": len(st.session_state.cdd_text)}
        )

        missing_titles = validate_approver_titles(st.session_state.excel_df, question_col, answer_col, qa_chain, session_logger)
        if missing_titles:
            st.markdown(
                "<div style='padding: 10px; margin: 10px 0;'>"
                "⚠️ Please ensure all approver job titles are included in the answers below."
                "</div>",
                unsafe_allow_html=True
            )
            styled_df = highlight_missing_titles(st.session_state.excel_df, missing_titles)
            excel_placeholder.dataframe(
                styled_df,
                use_container_width=True
            )
            for missing in missing_titles:
                st.markdown(
                    f"<span style='color: #00008B;'>Row {missing['index'] + 1}: Job title for {missing['approver_name']} not found in answer for question: {missing['question']}</span>",
                    unsafe_allow_html=True
                )
        else:
            st.session_state.cdd_validated = True
            session_logger.log(
                component="CDD Processing",
                message="Approver Title Validation",
                decision="Accepted",
                reason="All approver job titles found in answers",
                level="INFO",
                context={}
            )

        if invalid_answers:
            st.warning(f"⚠️ Invalid or empty answers found for questions: {', '.join(invalid_answers)}. Please edit manually before saving.")
            session_logger.log(
                component="CDD Processing",
                message="Answer Validation",
                decision="Warning",
                reason=f"Found {len(invalid_answers)} invalid answers",
                level="WARNING",
                context={"questions": ', '.join(invalid_answers)}
            )

        st.session_state.questions_processed = True
        session_logger.log(
            component="CDD Processing",
            message="Question Processing Completed",
            decision="Accepted",
            reason="All questions processed",
            level="INFO",
            context={}
        )

    # Always display the CDD DataFrame
    if not st.session_state.edit_mode:
        styled_df = highlight_missing_titles(st.session_state.excel_df, missing_titles)
        excel_placeholder.dataframe(
            styled_df,
            use_container_width=True
        )

    if not st.session_state.edit_mode and st.button("💾 Save Updated Excel", key="save_excel"):
        session_dir = os.path.join(SESSIONS_DIR, st.session_state.session_id)
        os.makedirs(session_dir, exist_ok=True)
        updated_excel_path = os.path.join(session_dir, "Control Design Document_updated.xlsx")
        try:
            st.session_state.excel_df.to_excel(updated_excel_path, engine="openpyxl", index=False)
            st.success(f"✅ Updated CDD saved to {updated_excel_path}")
            st.session_state.cdd_saved = True
            answers = st.session_state.excel_df[answer_col].dropna().astype(str)
            st.session_state.cdd_text = " ".join(answers)
            missing_titles = validate_approver_titles(st.session_state.excel_df, question_col, answer_col, qa_chain, session_logger)
            if not missing_titles:
                st.session_state.cdd_validated = True
                session_logger.log(
                    component="CDD Processing",
                    message="Approver Title Validation",
                    decision="Accepted",
                    reason="No missing approver titles after save",
                    level="INFO",
                    context={}
                )
            styled_df = highlight_missing_titles(st.session_state.excel_df, missing_titles)
            excel_placeholder.dataframe(
                styled_df,
                use_container_width=True
            )
            session_logger.log(
                component="CDD Processing",
                message="Excel Save",
                decision="Accepted",
                reason="Updated Excel saved successfully",
                level="INFO",
                context={"file": updated_excel_path}
            )
        except (IOError, PermissionError) as e:
            st.error(f"⚠️ Failed to save Excel: {str(e)}")
            session_logger.log(
                component="CDD Processing",
                message="Excel Save",
                decision="Rejected",
                reason=f"Failed to save Excel: {str(e)}",
                level="ERROR",
                context={"file": updated_excel_path}
            )

    doc = Document()
    doc.add_heading("Control Design Document Analysis Report", level=1)
    doc.add_paragraph(f"Control Owner: {st.session_state.pip_control_owner_name}")
    doc.add_paragraph(f"Designation: {st.session_state.pip_control_owner_title}")
    doc.add_paragraph()

    for idx, row in st.session_state.excel_df.iterrows():
        doc.add_heading(f"{row.name + 1}. {row[question_col]}", level=2)
        doc.add_paragraph(f"Answer: {row.get(answer_col, 'N/A')}")
        doc.add_paragraph()

    docx_filename = f"cdd_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        docx_path = save_docx(doc, docx_filename, session_logger)
        st.success(f"✅ Report generated: {docx_filename}")
        st.markdown(get_docx_download_link(docx_path, "Download Report"), unsafe_allow_html=True)
        session_logger.log(
            component="CDD Report",
            message=f"Generated DOCX: {docx_filename}",
            decision="Accepted",
            reason="DOCX report generated successfully",
            level="INFO",
            context={"file": docx_path}
        )
    except (IOError, PermissionError) as e:
        st.error(f"⚠️ Failed to generate DOCX report: {str(e)}")
        session_logger.log(
            component="CDD Report",
            message=f"Generating DOCX: {docx_filename}",
            decision="Rejected",
            reason=f"Failed to generate DOCX: {str(e)}",
            level="ERROR",
            context={"file": docx_path}
        )

    if st.session_state.questions_processed or qa_chain is None:
        if not st.session_state.cdd_saved or not st.session_state.cdd_validated:
            st.warning("⚠️ Please save the updated CDD and ensure all approver job titles are included before starting auditing.")
            session_logger.log(
                component="CDD Processing",
                message="Audit Start Check",
                decision="Rejected",
                reason="CDD not saved or approver titles not validated",
                level="WARNING",
                context={}
            )
        else:
            st.page_link("pages/audit_processing.py", label="Start Auditing")
            session_logger.log(
                component="CDD Processing",
                message="Audit Start Check",
                decision="Accepted",
                reason="CDD saved and validated, ready for auditing",
                level="INFO",
                context={}
            )

if __name__ == "__main__":
    process_cdd()